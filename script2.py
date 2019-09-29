from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import copy
import re

def find_occurances_in_paragraph(paragraph, search):
	return [m.start() for m in re.finditer(search, paragraph.text)]

def apply_format_to_range(paragraph, start, end, format_func):
	for run in get_target_runs(paragraph, start, end):
		format_func(run)

def get_target_runs(paragraph, start, end):
	targets = []

	#Must be done in a while loop because splitting the run will modify paragraph.runs
	i = 0
	past_start = False
	while(i < len(paragraph.runs)):
		run = paragraph.runs[i]
		run_start = sum([len(r.text) for r in paragraph.runs[:i]]) #inefficient but guaranteed correct
		run_end = run_start + len(run.text)
		
		run_contains_start = (run_start <= start <= run_end)
		run_contains_end = (run_start <= end <= run_end)

		#Split run in three, take middle part
		if(run_contains_start and run_contains_end):
			split_runs = split_run_in_three(paragraph, run, start-run_start, end-run_end)
			targets = [split_runs[1]]
			print([r.text for r in targets])
			return targets
		#Split run, take second half
		elif(run_contains_start and not run_contains_end):
			past_start = True
			split_runs = split_run_in_two(paragraph, run, start-run_start)
			targets.append(split_runs[1])
			i += 1 #skip run that was added by splitting run
		#Take whole run
		elif(past_start and not run_contains_end):
			targets.append(run)
		#Split run, take first half
		elif(past_start and run_contains_end):
			split_runs = split_run_in_two(paragraph, run, end-run_start)
			targets.append(split_runs[0])
			return targets
		i += 1
	return targets

def split_run_in_two(paragraph, run, split_index):
	index_in_paragraph = paragraph._p.index(run.element)

	text_before_split = run.text[0:split_index]
	text_after_split = run.text[split_index:]
	
	run.text = text_before_split
	new_run = paragraph.add_run(text_after_split)
	copy_format_manual(run, new_run)
	paragraph._p[index_in_paragraph+1:index_in_paragraph+1] = [new_run.element]
	return [run, new_run]

def split_run_in_three(paragraph, run, split_start, split_end):
	first_split = split_run_in_two(paragraph, run, split_end)
	second_split = split_run_in_two(paragraph, run, split_start)
	return second_split + [first_split[-1]]

def copy_format(runA, runB):
	#TODO: Find some way to copy runA font to runB
	#Does not work: 
	#runB.font = copy.deepcopy(runA.font)
	#runB.font.element.rPr = copy.deepcopy(runA.font.element.rPr)
	#runB.font.element = copy.deepcopy(runA.font.element)
	pass

def copy_format_manual(runA, runB):
	fontB = runB.font
	fontA = runA.font
	fontB.bold = fontA.bold
	fontB.italic = fontA.italic
	fontB.underline = fontA.underline
	fontB.strike = fontA.strike
	fontB.subscript = fontA.subscript
	fontB.superscript = fontA.superscript
	fontB.size = fontA.size
	fontB.highlight_color = fontA.highlight_color
	fontB.color.rgb = fontA.color.rgb
	#Probably others...

"""
The highlight_important function will take a text file and a docx file as inputs textFilename and docxFilename respectively. 
This file will have the same content as the input docx file but will have the important sentences highlited in yellow.
It also needs output document name as third input parameter.
It creates a new docx file with the outDocxFilename name as its result.
Remember to add proper extension to all filenames. 
"""    
def highlight_important(textFilename, inDocxFilename, outDocxFilename):
    textFile = open( textFilename , "r", encoding="utf8")
    sentences = textFile.readlines()
    doc = Document(inDocxFilename)
    print("Number of sentences to highlight: "+ str(len(sentences)) )
    for searches in sentences:
        if searches.endswith("\n"):
            search=searches[:-1]
        else:
            search=searches
        if search.endswith(" "):
            search=search[:-1]
        for paragraph in doc.paragraphs:
            format_func = lambda x:x.font.__setattr__('highlight_color', WD_COLOR_INDEX.YELLOW)
            for start in find_occurances_in_paragraph(paragraph, search):
                apply_format_to_range(paragraph, start, start + len(search), format_func)
    doc.save(outDocxFilename)
    print(sentences)
    textFile.close()

#Main function below
# highlight_important("out1.txt","Test.docx","Result.docx")
